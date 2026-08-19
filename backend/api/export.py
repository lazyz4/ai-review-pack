"""导出路由：POST /api/v1/export 与文件下载。

支持 Markdown / JSON / Word(docx) / PDF 四种格式；
Word 与 PDF 分别依赖 python-docx、reportlab，未安装时返回明确的 503 提示。
"""

from __future__ import annotations

import json
import os
from pathlib import Path

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse

from backend.schemas.course import ExportRequest, ExportResponse, GenerateAdvanceResponse
from backend.store import get_draft, get_request

router = APIRouter(prefix="/export", tags=["export"])

EXPORT_DIR = Path(os.getenv("EXPORT_DIR", "outputs/exports")).resolve()


@router.post("", response_model=ExportResponse, status_code=200, summary="导出最终复习包文件")
async def export_document(payload: ExportRequest) -> ExportResponse:
    """按指定格式生成复习包文件并返回下载信息。"""
    draft = get_draft(payload.outline_version)
    if draft is None:
        raise HTTPException(
            status_code=404,
            detail=f"未找到 outline_version={payload.outline_version}，请先生成复习包",
        )
    request = get_request(payload.outline_version)
    course_name = request.course_name if request else "软件测试"
    path = _write_export(draft, course_name, payload)
    return ExportResponse(
        file_name=path.name,
        file_size=path.stat().st_size,
        format=payload.format,
        download_url=f"/api/v1/export/download/{path.name}",
        outline_version=payload.outline_version,
    )


@router.get("/download/{file_name}", include_in_schema=False)
async def download_file(file_name: str) -> FileResponse:
    """返回已生成的导出文件（文件名做了防目录穿越处理）。"""
    safe_name = Path(file_name).name
    path = EXPORT_DIR / safe_name
    if not path.is_file():
        raise HTTPException(status_code=404, detail="导出文件不存在")
    return FileResponse(path, filename=path.name, media_type="application/octet-stream")


def _write_export(draft: GenerateAdvanceResponse, course_name: str, payload: ExportRequest) -> Path:
    """按请求格式写出文件并返回路径。"""
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    stem = f"{course_name}_{draft.outline_version.replace('.', '_')}_复习包"
    fmt = payload.format.upper()
    if fmt == "MARKDOWN":
        return _write_markdown(draft, course_name, payload, stem)
    if fmt == "JSON":
        return _write_json(draft, course_name, payload, stem)
    if fmt == "WORD":
        try:
            return _write_docx(draft, course_name, payload, stem)
        except ImportError as exc:
            raise HTTPException(status_code=503, detail="Word 导出需要安装 python-docx：pip install python-docx") from exc
    if fmt == "PDF":
        try:
            return _write_pdf(draft, course_name, payload, stem)
        except ImportError as exc:
            raise HTTPException(status_code=503, detail="PDF 导出需要安装 reportlab：pip install reportlab") from exc
    raise HTTPException(status_code=400, detail=f"不支持的导出格式：{payload.format}")


def _write_markdown(draft: GenerateAdvanceResponse, course_name: str, payload: ExportRequest, stem: str) -> Path:
    """生成 Markdown 导出文件。"""
    path = EXPORT_DIR / f"{stem}.md"
    lines = [
        f"# {course_name} 期末复习整合包",
        "",
        f"> 版本：{draft.outline_version}　生成时间：{draft.created_at:%Y-%m-%d %H:%M}",
        "",
        "## 一、执行摘要",
        "",
        draft.summary,
        "",
        "## 二、覆盖率",
        "",
        f"- 总知识点：{draft.coverage_metrics.total_topics}",
        f"- 已覆盖知识点：{draft.coverage_metrics.covered_topics}",
        f"- 覆盖率：{draft.coverage_metrics.coverage_rate:.0%}（目标 ≥ {draft.coverage_metrics.threshold:.0%}）",
        "",
        "## 三、知识点清单",
        "",
        "| 知识点ID | 名称 | 要点摘要 | 难度 | 高频 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for topic in draft.topics:
        lines.append(
            f"| {topic.topic_id} | {topic.name} | {topic.summary} | {topic.difficulty} | "
            f"{'是' if topic.is_high_frequency else '否'} |"
        )
    lines += ["", "## 四、题型模板与覆盖映射", "", "| 模板ID | 知识点ID | 题型 | 题干模板 | 难度 |", "| --- | --- | --- | --- | --- |"]
    for template in draft.templates:
        lines.append(f"| {template.template_id} | {template.topic_id} | {template.question_type} | {template.stem_template} | {template.difficulty} |")
    lines += ["", "## 五、示例题", ""]
    for question in draft.generated_questions:
        lines.append(f"### {question.question_id}（{question.question_type} / {question.difficulty}）")
        lines.append("")
        lines.append(f"**题干：** {question.stem}")
        if question.options:
            for option in question.options:
                lines.append(f"- {option}")
        lines.append("")
        lines.append(f"**答案要点：** {'；'.join(question.answer_points)}")
        lines.append("")
        lines.append(f"**解析要点：** {'；'.join(question.analysis_points)}")
        lines.append("")
    lines += ["## 六、个性化学习计划草案", "", "| 天 | 日期 | 阶段 | 目标知识点 | 每日题量 | 时长 |", "| --- | --- | --- | --- | --- | --- |"]
    for item in draft.study_plan_draft.items:
        lines.append(
            f"| 第{item.day}天 | {item.date or '-'} | {item.phase} | {'、'.join(item.focus_topics)} | "
            f"{item.daily_questions} | {item.duration_minutes}分钟 |"
        )
    lines += ["", "---", "", "> 本文件由《软件测试期末复习整合包》自动生成。"]
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def _write_json(draft: GenerateAdvanceResponse, course_name: str, payload: ExportRequest, stem: str) -> Path:
    """生成 JSON 导出文件（便于导入学习工具）。"""
    path = EXPORT_DIR / f"{stem}.json"
    data = {
        "outline_version": draft.outline_version,
        "course_name": course_name,
        "summary": draft.summary,
        "topics": [topic.model_dump(mode="json") for topic in draft.topics],
        "templates": [template.model_dump(mode="json") for template in draft.templates],
        "generated_questions": [question.model_dump(mode="json") for question in draft.generated_questions],
        "coverage_metrics": draft.coverage_metrics.model_dump(mode="json"),
        "study_plan_draft": draft.study_plan_draft.model_dump(mode="json"),
        "export_options": draft.export_options.model_dump(mode="json"),
        "created_at": draft.created_at.isoformat(),
    }
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _write_docx(draft: GenerateAdvanceResponse, course_name: str, payload: ExportRequest, stem: str) -> Path:
    """使用 python-docx 生成 Word 文档。"""
    from docx import Document

    path = EXPORT_DIR / f"{stem}.docx"
    document = Document()
    document.add_heading(f"{course_name} 期末复习整合包", 0)
    document.add_paragraph(f"版本：{draft.outline_version}　生成时间：{draft.created_at:%Y-%m-%d %H:%M}")

    document.add_heading("一、执行摘要", level=1)
    document.add_paragraph(draft.summary)

    document.add_heading("二、覆盖率", level=1)
    metrics = draft.coverage_metrics
    document.add_paragraph(
        f"总知识点：{metrics.total_topics}；已覆盖：{metrics.covered_topics}；"
        f"覆盖率：{metrics.coverage_rate:.0%}（目标 ≥ {metrics.threshold:.0%}）"
    )

    document.add_heading("三、知识点清单", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for index, title in enumerate(("知识点ID", "名称", "要点摘要", "难度", "是否高频")):
        table.rows[0].cells[index].text = title
    for topic in draft.topics:
        cells = table.add_row().cells
        cells[0].text = topic.topic_id
        cells[1].text = topic.name
        cells[2].text = topic.summary
        cells[3].text = topic.difficulty
        cells[4].text = "是" if topic.is_high_frequency else "否"

    document.add_heading("四、题型模板与覆盖映射", level=1)
    for template in draft.templates:
        document.add_paragraph(
            f"{template.template_id} | {template.topic_id} | {template.question_type} | "
            f"{template.stem_template} | {template.difficulty}"
        )

    document.add_heading("五、示例题", level=1)
    for question in draft.generated_questions:
        document.add_heading(f"{question.question_id}（{question.question_type} / {question.difficulty}）", level=2)
        document.add_paragraph(f"题干：{question.stem}")
        if question.options:
            for option in question.options:
                document.add_paragraph(option, style="List Bullet")
        document.add_paragraph(f"答案要点：{'；'.join(question.answer_points)}")
        document.add_paragraph(f"解析要点：{'；'.join(question.analysis_points)}")

    document.add_heading("六、个性化学习计划草案", level=1)
    plan_table = document.add_table(rows=1, cols=6)
    plan_table.style = "Table Grid"
    for index, title in enumerate(("天", "日期", "阶段", "目标知识点", "每日题量", "时长")):
        plan_table.rows[0].cells[index].text = title
    for item in draft.study_plan_draft.items:
        cells = plan_table.add_row().cells
        cells[0].text = str(item.day)
        cells[1].text = str(item.date or "")
        cells[2].text = item.phase
        cells[3].text = "、".join(item.focus_topics)
        cells[4].text = str(item.daily_questions)
        cells[5].text = f"{item.duration_minutes} 分钟"

    document.save(str(path))
    return path


def _write_pdf(draft: GenerateAdvanceResponse, course_name: str, payload: ExportRequest, stem: str) -> Path:
    """使用 reportlab 生成 PDF 文档（内置中文字体 STSong-Light）。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    font_name = "STSong-Light"
    title_style = ParagraphStyle("title", fontName=font_name, fontSize=18, leading=24, spaceAfter=8)
    heading_style = ParagraphStyle("heading", fontName=font_name, fontSize=14, leading=20, spaceBefore=10, spaceAfter=6)
    body_style = ParagraphStyle("body", fontName=font_name, fontSize=10.5, leading=16)

    path = EXPORT_DIR / f"{stem}.pdf"
    story = [
        Paragraph(f"{course_name} 期末复习整合包", title_style),
        Paragraph(f"版本：{draft.outline_version}　生成时间：{draft.created_at:%Y-%m-%d %H:%M}", body_style),
        Spacer(1, 12),
    ]
    story.append(Paragraph("一、执行摘要", heading_style))
    story.append(Paragraph(draft.summary, body_style))

    metrics = draft.coverage_metrics
    story.append(Paragraph("二、覆盖率", heading_style))
    story.append(
        Paragraph(
            f"总知识点：{metrics.total_topics}；已覆盖：{metrics.covered_topics}；"
            f"覆盖率：{metrics.coverage_rate:.0%}（目标 ≥ {metrics.threshold:.0%}）",
            body_style,
        )
    )

    story.append(Paragraph("三、知识点清单", heading_style))
    topic_rows = [["知识点ID", "名称", "要点摘要", "难度", "高频"]]
    for topic in draft.topics:
        topic_rows.append(
            [topic.topic_id, topic.name, topic.summary, topic.difficulty, "是" if topic.is_high_frequency else "否"]
        )
    topic_table = Table(topic_rows, colWidths=[60, 110, 220, 60, 50])
    topic_table.setStyle(
        TableStyle([("FONTNAME", (0, 0), (-1, -1), font_name), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey)])
    )
    story.append(topic_table)

    story.append(Paragraph("四、题型模板与覆盖映射", heading_style))
    for template in draft.templates:
        story.append(
            Paragraph(
                f"{template.template_id} | {template.topic_id} | {template.question_type} | "
                f"{template.stem_template} | {template.difficulty}",
                body_style,
            )
        )

    story.append(Paragraph("五、示例题", heading_style))
    for question in draft.generated_questions:
        story.append(Paragraph(f"{question.question_id}（{question.question_type} / {question.difficulty}）", heading_style))
        story.append(Paragraph(f"题干：{question.stem}", body_style))
        if question.options:
            for option in question.options:
                story.append(Paragraph(f"• {option}", body_style))
        story.append(Paragraph(f"答案要点：{'；'.join(question.answer_points)}", body_style))
        story.append(Paragraph(f"解析要点：{'；'.join(question.analysis_points)}", body_style))

    story.append(Paragraph("六、个性化学习计划草案", heading_style))
    plan_rows = [["天", "日期", "阶段", "目标知识点", "每日题量", "时长"]]
    for item in draft.study_plan_draft.items:
        plan_rows.append(
            [
                str(item.day),
                str(item.date or ""),
                item.phase,
                "、".join(item.focus_topics),
                str(item.daily_questions),
                f"{item.duration_minutes} 分钟",
            ]
        )
    plan_table = Table(plan_rows, colWidths=[35, 70, 70, 140, 60, 70])
    plan_table.setStyle(
        TableStyle([("FONTNAME", (0, 0), (-1, -1), font_name), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey)])
    )
    story.append(plan_table)

    document = SimpleDocTemplate(str(path), pagesize=A4)
    document.build(story)
    return path
